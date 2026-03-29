import streamlit as st
import fitz  # PyMuPDF
from docx import Document
import tempfile
import re

st.title("📄 PDF → Structured DOCX (Tamil + English)")

uploaded_file = st.file_uploader("Upload PDF", type=["pdf"])

def clean_text(text):
    return text.replace("\n", " ").strip()

def split_questions(text):
    # Split based on question numbers like "1." "2."
    return re.split(r"\n?\d+\.", text)

def extract_blocks(q):
    # Extract Tamil + English + options
    lines = q.split("\n")

    tamil_q = []
    eng_q = []
    options = []

    for line in lines:
        line = line.strip()

        if re.search(r"[அ-ஹ]", line):  # Tamil detection
            tamil_q.append(line)
        elif line.startswith(("A)", "B)", "C)", "D)")):
            options.append(line)
        else:
            eng_q.append(line)

    return " ".join(tamil_q), " ".join(eng_q), options

if uploaded_file:
    st.write("Processing...")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(uploaded_file.read())
        pdf_path = tmp.name

    pdf = fitz.open(pdf_path)

    full_text = ""

    for page in pdf:
        full_text += page.get_text()

    questions = split_questions(full_text)

    doc = Document()

    q_no = 1

    for q in questions:
        if len(q.strip()) < 10:
            continue

        tamil_q, eng_q, options = extract_blocks(q)

        # Tamil Question
        doc.add_paragraph(f"{q_no}. {tamil_q}")

        # Options
        for opt in options[:4]:
            doc.add_paragraph(opt)

        doc.add_paragraph("")  # spacing

        # English Question
        doc.add_paragraph(eng_q)

        for opt in options[:4]:
            doc.add_paragraph(opt)

        doc.add_paragraph("\n")

        q_no += 1

    output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(output_path)

    with open(output_path, "rb") as f:
        st.download_button(
            "Download DOCX",
            f.read(),
            file_name="structured_output.docx"
        )

    st.success("✅ Done")
